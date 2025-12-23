#!/usr/bin/env python3
"""
Invoice Expense Tracker with Dashboard & Zip Export
- Past year only
- OCR for scanned invoices
- Automatic expense totals (per-sender + per-currency)
- Interactive HTML dashboard
- Downloadable zip of all outputs
- Extensive logging and performance benchmarks
"""

import argparse
import base64
import csv
import hashlib
import json
import logging
import os
import re
import sys
import time
import zipfile
from collections import defaultdict
from contextlib import contextmanager
from datetime import date, timedelta, datetime
from email.utils import parseaddr
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple
from io import StringIO

from googleapiclient.discovery import build
from googleapiclient.errors import HttpError
from google.auth.transport.requests import Request
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import InstalledAppFlow

# Text extraction / OCR
import fitz  # PyMuPDF
from PIL import Image
import pytesseract
from pdf2image import convert_from_path
import docx

# ============== LOGGING SETUP ==============

class ColoredFormatter(logging.Formatter):
    """Custom formatter with colors for terminal output."""

    COLORS = {
        'DEBUG': '\033[36m',      # Cyan
        'INFO': '\033[32m',       # Green
        'WARNING': '\033[33m',    # Yellow
        'ERROR': '\033[31m',      # Red
        'CRITICAL': '\033[35m',   # Magenta
        'RESET': '\033[0m',       # Reset
        'BOLD': '\033[1m',        # Bold
        'DIM': '\033[2m',         # Dim
    }

    ICONS = {
        'DEBUG': '🔍',
        'INFO': '✅',
        'WARNING': '⚠️ ',
        'ERROR': '❌',
        'CRITICAL': '🔥',
    }

    def format(self, record):
        color = self.COLORS.get(record.levelname, self.COLORS['RESET'])
        reset = self.COLORS['RESET']
        dim = self.COLORS['DIM']
        icon = self.ICONS.get(record.levelname, '')

        # Format timestamp
        timestamp = datetime.fromtimestamp(record.created).strftime('%H:%M:%S.%f')[:-3]

        # Format the message
        formatted = f"{dim}[{timestamp}]{reset} {color}{icon} {record.levelname:<8}{reset} {record.getMessage()}"

        if record.exc_info:
            formatted += f"\n{self.formatException(record.exc_info)}"

        return formatted


def setup_logging(verbose: bool = False, log_file: Optional[Path] = None) -> logging.Logger:
    """Setup logging with colored console output and optional file output."""
    logger = logging.getLogger('invoice_tracker')
    logger.setLevel(logging.DEBUG if verbose else logging.INFO)
    logger.handlers.clear()

    # Console handler with colors
    console_handler = logging.StreamHandler(sys.stdout)
    console_handler.setLevel(logging.DEBUG if verbose else logging.INFO)
    console_handler.setFormatter(ColoredFormatter())
    logger.addHandler(console_handler)

    # File handler (no colors)
    if log_file:
        file_handler = logging.FileHandler(log_file, encoding='utf-8')
        file_handler.setLevel(logging.DEBUG)
        file_handler.setFormatter(logging.Formatter(
            '%(asctime)s | %(levelname)-8s | %(message)s',
            datefmt='%Y-%m-%d %H:%M:%S'
        ))
        logger.addHandler(file_handler)

    return logger


# ============== PERFORMANCE TRACKING ==============

class PerformanceTracker:
    """Track performance metrics throughout the script execution."""

    def __init__(self, logger: logging.Logger):
        self.logger = logger
        self.start_time = time.time()
        self.metrics = defaultdict(list)
        self.counters = defaultdict(int)
        self.current_timers = {}

    @contextmanager
    def timer(self, operation: str):
        """Context manager to time an operation."""
        start = time.time()
        self.logger.debug(f"⏱️  Starting: {operation}")
        try:
            yield
        finally:
            elapsed = time.time() - start
            self.metrics[operation].append(elapsed)
            self.logger.debug(f"⏱️  Completed: {operation} in {elapsed:.3f}s")

    def start_timer(self, name: str):
        """Start a named timer."""
        self.current_timers[name] = time.time()

    def stop_timer(self, name: str) -> float:
        """Stop a named timer and return elapsed time."""
        if name in self.current_timers:
            elapsed = time.time() - self.current_timers[name]
            self.metrics[name].append(elapsed)
            del self.current_timers[name]
            return elapsed
        return 0.0

    def increment(self, counter: str, amount: int = 1):
        """Increment a counter."""
        self.counters[counter] += amount

    def get_count(self, counter: str) -> int:
        """Get counter value."""
        return self.counters[counter]

    def log_progress(self, current: int, total: int, item_type: str = "items"):
        """Log progress with percentage and ETA."""
        if total == 0:
            return

        percent = (current / total) * 100
        elapsed = time.time() - self.start_time

        if current > 0:
            eta_seconds = (elapsed / current) * (total - current)
            eta_str = self._format_duration(eta_seconds)
        else:
            eta_str = "calculating..."

        bar_width = 30
        filled = int(bar_width * current / total)
        bar = '█' * filled + '░' * (bar_width - filled)

        self.logger.info(f"Progress: [{bar}] {current}/{total} {item_type} ({percent:.1f}%) | ETA: {eta_str}")

    def _format_duration(self, seconds: float) -> str:
        """Format duration in human readable format."""
        if seconds < 60:
            return f"{seconds:.1f}s"
        elif seconds < 3600:
            mins = int(seconds // 60)
            secs = int(seconds % 60)
            return f"{mins}m {secs}s"
        else:
            hours = int(seconds // 3600)
            mins = int((seconds % 3600) // 60)
            return f"{hours}h {mins}m"

    def print_summary(self):
        """Print a comprehensive performance summary."""
        total_time = time.time() - self.start_time

        self.logger.info("")
        self.logger.info("=" * 60)
        self.logger.info("📊 PERFORMANCE SUMMARY")
        self.logger.info("=" * 60)
        self.logger.info(f"Total execution time: {self._format_duration(total_time)}")
        self.logger.info("")

        # Counters
        if self.counters:
            self.logger.info("📈 COUNTERS:")
            for name, value in sorted(self.counters.items()):
                self.logger.info(f"   • {name}: {value:,}")
            self.logger.info("")

        # Timing metrics
        if self.metrics:
            self.logger.info("⏱️  TIMING BREAKDOWN:")
            for operation, times in sorted(self.metrics.items(), key=lambda x: -sum(x[1])):
                total = sum(times)
                avg = total / len(times) if times else 0
                count = len(times)
                pct = (total / total_time) * 100 if total_time > 0 else 0

                if count == 1:
                    self.logger.info(f"   • {operation}: {total:.3f}s ({pct:.1f}%)")
                else:
                    self.logger.info(f"   • {operation}: {total:.3f}s total, {avg:.3f}s avg, {count} calls ({pct:.1f}%)")
            self.logger.info("")

        self.logger.info("=" * 60)


# ============== CONSTANTS ==============

SCOPES = ["https://www.googleapis.com/auth/gmail.readonly"]

ALLOWED_EXTS = {
    ".pdf", ".docx",
    ".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp", ".webp"
}

DEFAULT_KEYWORDS = [
    "invoice", "invoices", "receipt", "receipts", "tax invoice", "bill",
    "חשבונית", "חשבוניות", "קבלה", "קבלות", "חשבונית מס", "חשבונית מס קבלה", "סכום לתשלום"
]

TOTAL_KEYWORDS = [
    "total", "grand total", "amount due", "total due", "balance due",
    "סה\"כ", "סהכ", "סכום כולל", "סכום לתשלום", "לתשלום", "סהכ לתשלום"
]

TAX_WORDS = ["vat", "tax", "מע\"מ", "מעמ", "מס"]

URL_REGEX = re.compile(r"""(?xi)\b(https?://[^\s<>"'\]]+|www\.[^\s<>"'\]]+)\b""")

AMOUNT_REGEX = re.compile(
    r"(?<!\w)(?:₪|\$|€|£)?\s*(?:\d{1,3}(?:[.,\s]\d{3})+|\d+)(?:[.,]\d{2})?\s*(?:₪|ש\"ח|שח|NIS|ILS|USD|EUR|GBP|\$|€|£)?"
)

CURRENCY_SYMBOLS = {
    "ILS": "₪",
    "USD": "$",
    "EUR": "€",
    "GBP": "£",
    "CAD": "C$",
    "AUD": "A$",
    "JPY": "¥",
    "CHF": "Fr.",
    "CNY": "¥",
    "INR": "₹",
    "BRL": "R$",
    "MXN": "MX$",
    "SGD": "S$",
    "HKD": "HK$",
    "KRW": "₩",
    "RUB": "₽",
    "TRY": "₺",
    "PLN": "zł",
    "THB": "฿",
    "ZAR": "R",
    "SEK": "kr",
    "NOK": "kr",
    "DKK": "kr",
    "NZD": "NZ$",
    "AED": "د.إ",
    "UNK": "?"
}


# ============== UTILITY FUNCTIONS ==============

def ensure_dir(p: Path) -> None:
    p.mkdir(parents=True, exist_ok=True)


def sanitize_filename(name: str) -> str:
    name = (name or "").strip().replace("\0", "")
    return re.sub(r"[^\w\-.() \[\]]+", "_", name) or "unknown"


def decode_b64(data: str) -> bytes:
    missing_padding = len(data) % 4
    if missing_padding:
        data += "=" * (4 - missing_padding)
    return base64.urlsafe_b64decode(data.encode("utf-8"))


def format_bytes(size: int) -> str:
    """Format bytes to human readable string."""
    for unit in ['B', 'KB', 'MB', 'GB']:
        if size < 1024:
            return f"{size:.1f}{unit}"
        size /= 1024
    return f"{size:.1f}TB"


# ============== GMAIL AUTH ==============

def load_or_auth(creds_path: Path, token_path: Path, logger: logging.Logger, perf: PerformanceTracker) -> Credentials:
    logger.info(f"🔐 Checking authentication...")
    logger.debug(f"   Credentials file: {creds_path}")
    logger.debug(f"   Token file: {token_path}")

    creds: Optional[Credentials] = None

    with perf.timer("authentication"):
        if token_path.exists():
            logger.info(f"   Found existing token, loading...")
            creds = Credentials.from_authorized_user_file(str(token_path), SCOPES)
            logger.debug(f"   Token valid: {creds.valid if creds else False}")
            logger.debug(f"   Token expired: {creds.expired if creds else 'N/A'}")

        if not creds or not creds.valid:
            if creds and creds.expired and creds.refresh_token:
                logger.info(f"   Token expired, refreshing...")
                creds.refresh(Request())
                logger.info(f"   Token refreshed successfully")
            else:
                logger.warning(f"   No valid token found, starting OAuth flow...")
                logger.info(f"   🌐 Opening browser for authentication...")
                flow = InstalledAppFlow.from_client_secrets_file(str(creds_path), SCOPES)
                creds = flow.run_local_server(port=0)
                logger.info(f"   Authentication successful!")

            ensure_dir(token_path.parent)
            token_path.write_text(creds.to_json(), encoding="utf-8")
            logger.debug(f"   Token saved to {token_path}")

    logger.info(f"✅ Authentication complete")
    return creds


def gmail_service(creds: Credentials):
    return build("gmail", "v1", credentials=creds, cache_discovery=False)


# ============== GMAIL OPERATIONS ==============

def build_gmail_query(keywords: List[str], after: str) -> str:
    kw = " OR ".join([f'"{k}"' if " " in k else k for k in keywords])
    attachment_part = "(has:attachment OR filename:pdf OR filename:docx OR filename:jpg OR filename:png)"
    return f"({attachment_part}) ({kw}) after:{after}"


def list_messages(service, query: str, max_messages: int, logger: logging.Logger, perf: PerformanceTracker) -> List[str]:
    logger.info(f"📬 Fetching message list from Gmail...")
    logger.debug(f"   Query: {query}")
    logger.debug(f"   Max messages: {max_messages}")

    ids: List[str] = []
    page_token = None
    page_count = 0

    with perf.timer("list_messages"):
        while True:
            page_count += 1
            logger.debug(f"   Fetching page {page_count}...")

            resp = service.users().messages().list(
                userId="me", q=query, pageToken=page_token, maxResults=min(500, max_messages - len(ids))
            ).execute()

            msgs = resp.get("messages", [])
            ids.extend([m["id"] for m in msgs])

            logger.debug(f"   Page {page_count}: got {len(msgs)} messages (total: {len(ids)})")

            page_token = resp.get("nextPageToken")
            if not page_token or len(ids) >= max_messages:
                break

    perf.increment("messages_found", len(ids[:max_messages]))
    logger.info(f"✅ Found {len(ids[:max_messages])} messages matching query")
    return ids[:max_messages]


def get_header(headers: List[Dict[str, str]], name: str) -> str:
    for h in headers:
        if h.get("name", "").lower() == name.lower():
            return h.get("value", "")
    return ""


def extract_text_parts(payload: Dict[str, Any]) -> Tuple[str, str]:
    plain_chunks: List[str] = []
    html_chunks: List[str] = []

    def walk(part: Dict[str, Any]):
        mime = part.get("mimeType", "")
        body = part.get("body", {}) or {}
        data = body.get("data")
        if data and mime in ("text/plain", "text/html"):
            try:
                txt = decode_b64(data).decode("utf-8", errors="replace")
            except Exception:
                txt = ""
            (plain_chunks if mime == "text/plain" else html_chunks).append(txt)

        for sub in (part.get("parts") or []):
            walk(sub)

    walk(payload)
    return ("\n".join(plain_chunks), "\n".join(html_chunks))


def extract_links(text: str) -> List[str]:
    links = []
    for m in URL_REGEX.finditer(text or ""):
        url = m.group(1)
        if url.startswith("www."):
            url = "http://" + url
        links.append(url.rstrip(").,;!\"'<>]"))
    seen = set()
    out = []
    for u in links:
        if u not in seen:
            out.append(u)
            seen.add(u)
    return out


def iter_attachments(payload: Dict[str, Any]) -> List[Dict[str, Any]]:
    atts: List[Dict[str, Any]] = []

    def walk(part: Dict[str, Any]):
        filename = part.get("filename") or ""
        body = part.get("body", {}) or {}
        att_id = body.get("attachmentId")
        if filename and att_id:
            atts.append({
                "filename": filename,
                "mimeType": part.get("mimeType", ""),
                "attachmentId": att_id,
                "size": body.get("size", 0),
            })
        for sub in (part.get("parts") or []):
            walk(sub)

    walk(payload)
    return atts


def download_attachment(service, msg_id: str, attachment_id: str) -> bytes:
    att = service.users().messages().attachments().get(
        userId="me", messageId=msg_id, id=attachment_id
    ).execute()
    return decode_b64(att.get("data", ""))


def sha256_bytes(b: bytes) -> str:
    return hashlib.sha256(b).hexdigest()


def parse_gmail_internal_date(ms: str) -> str:
    try:
        dt = datetime.utcfromtimestamp(int(ms) / 1000)
        return dt.isoformat() + "Z"
    except Exception:
        return ""


# ============== AMOUNT EXTRACTION ==============

def normalize_amount_str(s: str) -> Optional[float]:
    s = s.strip()
    s = re.sub(r"[^\d,.\s]", "", s)
    s = re.sub(r"\s+", "", s)

    if not s or not re.search(r"\d", s):
        return None

    has_comma = "," in s
    has_dot = "." in s

    if has_comma and has_dot:
        if s.rfind(",") > s.rfind("."):
            s = s.replace(".", "")
            s = s.replace(",", ".")
        else:
            s = s.replace(",", "")
    elif has_comma and not has_dot:
        if re.search(r",\d{2}$", s):
            s = s.replace(",", ".")
        else:
            s = s.replace(",", "")
    else:
        if has_dot and not re.search(r"\.\d{2}$", s):
            s = s.replace(".", "")

    try:
        return float(s)
    except Exception:
        return None


def detect_currency(context: str) -> str:
    """
    Detect currency from context around an amount.
    Returns currency code: ILS, USD, EUR, GBP, CAD, AUD, JPY, CHF, CNY, INR, BRL, MXN, or UNK
    """
    ctx = context.lower()
    original = context  # Keep original for symbol matching

    # Israeli Shekel (ILS) - check first as it's common in Hebrew invoices
    ils_patterns = [
        "₪", "ש\"ח", "שח", "ש'ח", "שקל", "שקלים", "שקל חדש", "שקלים חדשים",
        "nis", "ils", "shekel", "shekels", "new israeli shekel", "israeli shekel",
        "ש״ח",  # Hebrew quotation mark variant
    ]
    for pat in ils_patterns:
        if pat in ctx or pat in original:
            return "ILS"

    # US Dollar (USD)
    usd_patterns = [
        "usd", "us$", "u.s. dollar", "us dollar", "u.s.d", "american dollar",
        "united states dollar", "dollar", "dollars",
    ]
    if "$" in original and "ca$" not in ctx and "a$" not in ctx and "hk$" not in ctx and "s$" not in ctx:
        return "USD"
    for pat in usd_patterns:
        if pat in ctx:
            return "USD"

    # Euro (EUR)
    eur_patterns = [
        "€", "eur", "euro", "euros", "€ur",
    ]
    for pat in eur_patterns:
        if pat in ctx or pat in original:
            return "EUR"

    # British Pound (GBP)
    gbp_patterns = [
        "£", "gbp", "pound", "pounds", "sterling", "british pound", "pound sterling", "quid",
    ]
    for pat in gbp_patterns:
        if pat in ctx or pat in original:
            return "GBP"

    # Canadian Dollar (CAD)
    cad_patterns = ["cad", "ca$", "c$", "canadian dollar", "canadian dollars"]
    for pat in cad_patterns:
        if pat in ctx:
            return "CAD"

    # Australian Dollar (AUD)
    aud_patterns = ["aud", "a$", "au$", "australian dollar", "australian dollars"]
    for pat in aud_patterns:
        if pat in ctx:
            return "AUD"

    # Japanese Yen (JPY)
    jpy_patterns = ["¥", "jpy", "yen", "円", "japanese yen"]
    for pat in jpy_patterns:
        if pat in ctx or pat in original:
            return "JPY"

    # Swiss Franc (CHF)
    chf_patterns = ["chf", "sfr", "swiss franc", "swiss francs", "franken"]
    for pat in chf_patterns:
        if pat in ctx:
            return "CHF"

    # Chinese Yuan (CNY)
    cny_patterns = ["cny", "rmb", "yuan", "renminbi", "元", "人民币"]
    for pat in cny_patterns:
        if pat in ctx or pat in original:
            return "CNY"

    # Indian Rupee (INR)
    inr_patterns = ["₹", "inr", "rupee", "rupees", "indian rupee", "rs", "rs."]
    for pat in inr_patterns:
        if pat in ctx or pat in original:
            return "INR"

    # Brazilian Real (BRL)
    brl_patterns = ["r$", "brl", "real", "reais", "brazilian real"]
    for pat in brl_patterns:
        if pat in ctx:
            return "BRL"

    # Mexican Peso (MXN)
    mxn_patterns = ["mxn", "mx$", "mexican peso", "pesos mexicanos"]
    for pat in mxn_patterns:
        if pat in ctx:
            return "MXN"

    # Singapore Dollar (SGD)
    sgd_patterns = ["sgd", "s$", "singapore dollar"]
    for pat in sgd_patterns:
        if pat in ctx:
            return "SGD"

    # Hong Kong Dollar (HKD)
    hkd_patterns = ["hkd", "hk$", "hong kong dollar"]
    for pat in hkd_patterns:
        if pat in ctx:
            return "HKD"

    # South Korean Won (KRW)
    krw_patterns = ["₩", "krw", "won", "원", "korean won"]
    for pat in krw_patterns:
        if pat in ctx or pat in original:
            return "KRW"

    # Russian Ruble (RUB)
    rub_patterns = ["₽", "rub", "ruble", "rubles", "рубль", "рублей", "руб"]
    for pat in rub_patterns:
        if pat in ctx or pat in original:
            return "RUB"

    # Turkish Lira (TRY)
    try_patterns = ["₺", "try", "tl", "turkish lira", "lira"]
    for pat in try_patterns:
        if pat in ctx or pat in original:
            return "TRY"

    # Polish Zloty (PLN)
    pln_patterns = ["zł", "pln", "zloty", "złoty", "złotych"]
    for pat in pln_patterns:
        if pat in ctx or pat in original:
            return "PLN"

    # Thai Baht (THB)
    thb_patterns = ["฿", "thb", "baht"]
    for pat in thb_patterns:
        if pat in ctx or pat in original:
            return "THB"

    # South African Rand (ZAR)
    zar_patterns = ["zar", "rand", "south african rand"]
    for pat in zar_patterns:
        if pat in ctx:
            return "ZAR"

    # Swedish Krona (SEK)
    sek_patterns = ["sek", "kr", "krona", "kronor", "swedish krona"]
    for pat in sek_patterns:
        if pat in ctx:
            return "SEK"

    # Norwegian Krone (NOK)
    nok_patterns = ["nok", "norwegian krone", "norske kroner"]
    for pat in nok_patterns:
        if pat in ctx:
            return "NOK"

    # Danish Krone (DKK)
    dkk_patterns = ["dkk", "danish krone", "danske kroner"]
    for pat in dkk_patterns:
        if pat in ctx:
            return "DKK"

    # New Zealand Dollar (NZD)
    nzd_patterns = ["nzd", "nz$", "new zealand dollar"]
    for pat in nzd_patterns:
        if pat in ctx:
            return "NZD"

    # UAE Dirham (AED)
    aed_patterns = ["aed", "dirham", "dirhams", "uae dirham", "د.إ"]
    for pat in aed_patterns:
        if pat in ctx or pat in original:
            return "AED"

    return "UNK"


def extract_best_total(text: str, logger: logging.Logger) -> Optional[Dict[str, Any]]:
    if not text or len(text.strip()) < 10:
        return None

    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    candidates: List[Dict[str, Any]] = []

    def add_candidate(amount_str: str, context: str, score: int, source: str):
        val = normalize_amount_str(amount_str)
        if val is None:
            return
        curr = detect_currency(context)
        candidates.append({
            "amount": val,
            "currency": curr,
            "context": context[:200],
            "score": score,
            "source": source,
        })

    for idx, ln in enumerate(lines):
        low = ln.lower()
        if any(k in low for k in TOTAL_KEYWORDS):
            score = 100
            if any(tw in low for tw in TAX_WORDS):
                score -= 30
            for m in AMOUNT_REGEX.finditer(ln):
                add_candidate(m.group(0), ln, score, "keyword_line")
            for j in range(1, 3):
                if idx + j < len(lines):
                    near = lines[idx + j]
                    near_low = near.lower()
                    near_score = 75
                    if any(tw in near_low for tw in TAX_WORDS):
                        near_score -= 25
                    for m in AMOUNT_REGEX.finditer(near):
                        add_candidate(m.group(0), ln + " || " + near, near_score, "near_keyword")

    pat = re.compile(r"(?ix)(total|סה\"כ|סהכ|סכום לתשלום|amount due).{0,40}?(\d[\d,.\s]{1,20}(?:[.,]\d{2})?)")
    for m in pat.finditer(text):
        ctx = text[max(0, m.start()-60): m.end()+60]
        add_candidate(m.group(2), ctx, 60, "global_pattern")

    if not candidates:
        logger.debug(f"      No amount candidates found in text ({len(text)} chars)")
        return None

    candidates.sort(key=lambda c: (c["currency"] != "UNK", c["score"], c["amount"]), reverse=True)
    best = candidates[0]
    logger.debug(f"      Found {len(candidates)} candidates, best: {best['currency']} {best['amount']:.2f} (score: {best['score']}, method: {best['source']})")
    return best


# ============== OCR & TEXT EXTRACTION ==============

def extract_text_from_pdf(pdf_path: Path, ocr: bool, ocr_max_pages: int, logger: logging.Logger, perf: PerformanceTracker) -> str:
    txt = ""

    # Try native text first
    with perf.timer("pdf_native_extraction"):
        try:
            doc = fitz.open(str(pdf_path))
            chunks = []
            page_count = len(doc)
            for page in doc:
                chunks.append(page.get_text("text"))
            txt = "\n".join(chunks).strip()
            doc.close()
            logger.debug(f"      PDF native text: {len(txt)} chars from {page_count} pages")
        except Exception as e:
            logger.warning(f"      PDF native extraction failed: {e}")
            txt = ""

    # If text is too small, OCR it (scanned pdf)
    if ocr and len(txt) < 200:
        logger.debug(f"      Text too short ({len(txt)} chars), attempting OCR...")
        perf.increment("ocr_attempts")

        with perf.timer("pdf_ocr"):
            try:
                images = convert_from_path(str(pdf_path), first_page=1, last_page=min(ocr_max_pages, 50))
                logger.debug(f"      Converted {len(images)} pages to images for OCR")

                ocr_chunks = []
                for i, img in enumerate(images[:ocr_max_pages]):
                    logger.debug(f"      OCR processing page {i+1}/{len(images)}...")
                    ocr_chunks.append(pytesseract.image_to_string(img))

                ocr_text = "\n".join(ocr_chunks)
                txt = (txt + "\n\n" + ocr_text).strip()
                logger.debug(f"      OCR extracted: {len(ocr_text)} chars")
                perf.increment("ocr_success")
            except Exception as e:
                logger.warning(f"      OCR failed: {e}")
                perf.increment("ocr_failed")

    return txt


def extract_text_from_image(img_path: Path, logger: logging.Logger, perf: PerformanceTracker) -> str:
    logger.debug(f"      Extracting text from image via OCR...")
    perf.increment("image_ocr_attempts")

    with perf.timer("image_ocr"):
        try:
            img = Image.open(img_path)
            text = pytesseract.image_to_string(img).strip()
            logger.debug(f"      Image OCR extracted: {len(text)} chars")
            perf.increment("image_ocr_success")
            return text
        except Exception as e:
            logger.warning(f"      Image OCR failed: {e}")
            perf.increment("image_ocr_failed")
            return ""


def extract_text_from_docx(docx_path: Path, logger: logging.Logger, perf: PerformanceTracker) -> str:
    logger.debug(f"      Extracting text from DOCX...")

    with perf.timer("docx_extraction"):
        try:
            d = docx.Document(str(docx_path))
            text = "\n".join(p.text for p in d.paragraphs).strip()
            logger.debug(f"      DOCX extracted: {len(text)} chars from {len(d.paragraphs)} paragraphs")
            return text
        except Exception as e:
            logger.warning(f"      DOCX extraction failed: {e}")
            return ""


def analyze_file(path: Path, enable_ocr: bool, ocr_max_pages: int, logger: logging.Logger, perf: PerformanceTracker) -> Dict[str, Any]:
    ext = path.suffix.lower()
    text = ""

    logger.debug(f"      Analyzing file: {path.name} ({ext})")

    with perf.timer("file_analysis"):
        if ext == ".pdf":
            text = extract_text_from_pdf(path, enable_ocr, ocr_max_pages, logger, perf)
            perf.increment("pdfs_processed")
        elif ext == ".docx":
            text = extract_text_from_docx(path, logger, perf)
            perf.increment("docx_processed")
        elif ext in {".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp", ".webp"}:
            if enable_ocr:
                text = extract_text_from_image(path, logger, perf)
            perf.increment("images_processed")

        best = extract_best_total(text, logger)

    return {
        "extracted_text_len": len(text),
        "best_total": best,
    }


# ============== DASHBOARD GENERATION ==============

def generate_dashboard_html(
    account_label: str,
    file_rows: List[Dict],
    total_by_currency: Dict[str, float],
    total_by_sender_currency: Dict[Tuple[str, str], float],
    generated_at: str,
    perf_summary: Dict[str, Any]
) -> str:
    """Generate an interactive HTML dashboard."""

    # Prepare data for charts
    currency_data = [{"currency": k, "amount": v} for k, v in sorted(total_by_currency.items(), key=lambda x: -x[1])]

    # Top 10 senders by total amount
    sender_totals = defaultdict(float)
    for (sender, curr), amt in total_by_sender_currency.items():
        sender_totals[sender] += amt
    top_senders = sorted(sender_totals.items(), key=lambda x: -x[1])[:10]

    # Monthly breakdown
    monthly_totals = defaultdict(lambda: defaultdict(float))
    for row in file_rows:
        try:
            dt = datetime.fromisoformat(row["date_utc"].replace("Z", "+00:00"))
            month_key = dt.strftime("%Y-%m")
            monthly_totals[month_key][row["currency"]] += float(row["amount"])
        except:
            pass

    months_sorted = sorted(monthly_totals.keys())

    html = f'''<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Invoice Expenses Dashboard - {account_label}</title>
    <script src="https://cdn.jsdelivr.net/npm/chart.js@4.4.1/dist/chart.umd.min.js"></script>
    <style>
        :root {{
            --bg-primary: #0a0a0f;
            --bg-secondary: #12121a;
            --bg-card: #1a1a24;
            --text-primary: #ffffff;
            --text-secondary: #a0a0b0;
            --accent: #6366f1;
            --accent-hover: #818cf8;
            --success: #22c55e;
            --warning: #f59e0b;
            --danger: #ef4444;
            --border: #2a2a3a;
        }}

        * {{
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }}

        body {{
            font-family: 'Inter', -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
            background: var(--bg-primary);
            color: var(--text-primary);
            min-height: 100vh;
            line-height: 1.6;
        }}

        .container {{
            max-width: 1400px;
            margin: 0 auto;
            padding: 2rem;
        }}

        header {{
            display: flex;
            justify-content: space-between;
            align-items: center;
            margin-bottom: 2rem;
            padding-bottom: 1.5rem;
            border-bottom: 1px solid var(--border);
        }}

        h1 {{
            font-size: 2rem;
            font-weight: 700;
            background: linear-gradient(135deg, var(--accent), #a855f7);
            -webkit-background-clip: text;
            -webkit-text-fill-color: transparent;
            background-clip: text;
        }}

        .meta {{
            color: var(--text-secondary);
            font-size: 0.875rem;
        }}

        .stats-grid {{
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(200px, 1fr));
            gap: 1.5rem;
            margin-bottom: 2rem;
        }}

        .stat-card {{
            background: var(--bg-card);
            border: 1px solid var(--border);
            border-radius: 12px;
            padding: 1.5rem;
            transition: transform 0.2s, box-shadow 0.2s;
        }}

        .stat-card:hover {{
            transform: translateY(-2px);
            box-shadow: 0 8px 30px rgba(99, 102, 241, 0.15);
        }}

        .stat-label {{
            font-size: 0.875rem;
            color: var(--text-secondary);
            margin-bottom: 0.5rem;
            text-transform: uppercase;
            letter-spacing: 0.05em;
        }}

        .stat-value {{
            font-size: 1.75rem;
            font-weight: 700;
        }}

        .stat-value.ils {{ color: #22c55e; }}
        .stat-value.usd {{ color: #3b82f6; }}
        .stat-value.eur {{ color: #f59e0b; }}
        .stat-value.gbp {{ color: #ec4899; }}

        .charts-grid {{
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(400px, 1fr));
            gap: 1.5rem;
            margin-bottom: 2rem;
        }}

        .chart-card {{
            background: var(--bg-card);
            border: 1px solid var(--border);
            border-radius: 12px;
            padding: 1.5rem;
        }}

        .chart-title {{
            font-size: 1.125rem;
            font-weight: 600;
            margin-bottom: 1rem;
            color: var(--text-primary);
        }}

        .chart-container {{
            position: relative;
            height: 300px;
        }}

        .table-section {{
            background: var(--bg-card);
            border: 1px solid var(--border);
            border-radius: 12px;
            overflow: hidden;
            margin-bottom: 2rem;
        }}

        .table-header {{
            display: flex;
            justify-content: space-between;
            align-items: center;
            padding: 1.25rem 1.5rem;
            border-bottom: 1px solid var(--border);
        }}

        .table-title {{
            font-size: 1.125rem;
            font-weight: 600;
        }}

        .search-box {{
            background: var(--bg-secondary);
            border: 1px solid var(--border);
            border-radius: 8px;
            padding: 0.5rem 1rem;
            color: var(--text-primary);
            font-size: 0.875rem;
            width: 250px;
            outline: none;
            transition: border-color 0.2s;
        }}

        .search-box:focus {{
            border-color: var(--accent);
        }}

        .table-wrapper {{
            overflow-x: auto;
        }}

        table {{
            width: 100%;
            border-collapse: collapse;
        }}

        th, td {{
            padding: 1rem 1.25rem;
            text-align: left;
            border-bottom: 1px solid var(--border);
        }}

        th {{
            background: var(--bg-secondary);
            font-weight: 600;
            font-size: 0.75rem;
            text-transform: uppercase;
            letter-spacing: 0.05em;
            color: var(--text-secondary);
            cursor: pointer;
            user-select: none;
        }}

        th:hover {{
            color: var(--accent);
        }}

        tr:hover td {{
            background: var(--bg-secondary);
        }}

        .amount {{
            font-family: 'JetBrains Mono', monospace;
            font-weight: 600;
        }}

        .currency-badge {{
            display: inline-block;
            padding: 0.25rem 0.5rem;
            border-radius: 4px;
            font-size: 0.75rem;
            font-weight: 600;
        }}

        .currency-badge.ILS {{ background: rgba(34, 197, 94, 0.2); color: #22c55e; }}
        .currency-badge.USD {{ background: rgba(59, 130, 246, 0.2); color: #3b82f6; }}
        .currency-badge.EUR {{ background: rgba(245, 158, 11, 0.2); color: #f59e0b; }}
        .currency-badge.GBP {{ background: rgba(236, 72, 153, 0.2); color: #ec4899; }}
        .currency-badge.UNK {{ background: rgba(160, 160, 176, 0.2); color: #a0a0b0; }}

        .download-btn {{
            display: inline-flex;
            align-items: center;
            gap: 0.5rem;
            background: var(--accent);
            color: white;
            padding: 0.75rem 1.5rem;
            border-radius: 8px;
            text-decoration: none;
            font-weight: 600;
            transition: background 0.2s;
            border: none;
            cursor: pointer;
            font-size: 0.875rem;
        }}

        .download-btn:hover {{
            background: var(--accent-hover);
        }}

        .actions {{
            display: flex;
            gap: 1rem;
        }}

        .empty-state {{
            text-align: center;
            padding: 3rem;
            color: var(--text-secondary);
        }}

        .truncate {{
            max-width: 200px;
            overflow: hidden;
            text-overflow: ellipsis;
            white-space: nowrap;
        }}

        .perf-section {{
            background: var(--bg-card);
            border: 1px solid var(--border);
            border-radius: 12px;
            padding: 1.5rem;
            margin-bottom: 2rem;
        }}

        .perf-grid {{
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(150px, 1fr));
            gap: 1rem;
            margin-top: 1rem;
        }}

        .perf-item {{
            text-align: center;
        }}

        .perf-value {{
            font-size: 1.5rem;
            font-weight: 700;
            color: var(--accent);
        }}

        .perf-label {{
            font-size: 0.75rem;
            color: var(--text-secondary);
            text-transform: uppercase;
        }}

        @media (max-width: 768px) {{
            .container {{
                padding: 1rem;
            }}

            header {{
                flex-direction: column;
                gap: 1rem;
                text-align: center;
            }}

            .charts-grid {{
                grid-template-columns: 1fr;
            }}

            .table-header {{
                flex-direction: column;
                gap: 1rem;
            }}

            .search-box {{
                width: 100%;
            }}
        }}
    </style>
</head>
<body>
    <div class="container">
        <header>
            <div>
                <h1>Invoice Expenses Dashboard</h1>
                <p class="meta">Account: {account_label} | Generated: {generated_at}</p>
            </div>
            <div class="actions">
                <button class="download-btn" onclick="exportCSV()">
                    <svg width="16" height="16" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2">
                        <path d="M21 15v4a2 2 0 0 1-2 2H5a2 2 0 0 1-2-2v-4"/>
                        <polyline points="7 10 12 15 17 10"/>
                        <line x1="12" y1="15" x2="12" y2="3"/>
                    </svg>
                    Export CSV
                </button>
            </div>
        </header>

        <div class="perf-section">
            <h3 class="chart-title">Processing Statistics</h3>
            <div class="perf-grid">
                <div class="perf-item">
                    <div class="perf-value">{perf_summary.get('messages_processed', 0)}</div>
                    <div class="perf-label">Emails Scanned</div>
                </div>
                <div class="perf-item">
                    <div class="perf-value">{perf_summary.get('attachments_downloaded', 0)}</div>
                    <div class="perf-label">Attachments</div>
                </div>
                <div class="perf-item">
                    <div class="perf-value">{perf_summary.get('total_bytes', '0 B')}</div>
                    <div class="perf-label">Data Downloaded</div>
                </div>
                <div class="perf-item">
                    <div class="perf-value">{perf_summary.get('ocr_performed', 0)}</div>
                    <div class="perf-label">OCR Scans</div>
                </div>
                <div class="perf-item">
                    <div class="perf-value">{perf_summary.get('processing_time', '0s')}</div>
                    <div class="perf-label">Processing Time</div>
                </div>
            </div>
        </div>

        <div class="stats-grid">
            {''.join(f"""
            <div class="stat-card">
                <div class="stat-label">Total {item['currency']}</div>
                <div class="stat-value {item['currency'].lower()}">{CURRENCY_SYMBOLS.get(item['currency'], '')}{item['amount']:,.2f}</div>
            </div>
            """ for item in sorted(currency_data, key=lambda x: -x['amount']) if item['amount'] > 0) if currency_data else '<div class="stat-card"><div class="stat-label">No expenses found</div><div class="stat-value">-</div></div>'}
            <div class="stat-card">
                <div class="stat-label">Total Invoices</div>
                <div class="stat-value">{len(file_rows)}</div>
            </div>
            <div class="stat-card">
                <div class="stat-label">Unique Senders</div>
                <div class="stat-value">{len(sender_totals)}</div>
            </div>
        </div>

        <div class="charts-grid">
            <div class="chart-card">
                <h3 class="chart-title">Expenses by Currency</h3>
                <div class="chart-container">
                    <canvas id="currencyChart"></canvas>
                </div>
            </div>
            <div class="chart-card">
                <h3 class="chart-title">Top 10 Senders</h3>
                <div class="chart-container">
                    <canvas id="sendersChart"></canvas>
                </div>
            </div>
            <div class="chart-card" style="grid-column: 1 / -1;">
                <h3 class="chart-title">Monthly Trend</h3>
                <div class="chart-container" style="height: 250px;">
                    <canvas id="monthlyChart"></canvas>
                </div>
            </div>
        </div>

        <div class="table-section">
            <div class="table-header">
                <h3 class="table-title">All Invoices ({len(file_rows)} total)</h3>
                <input type="text" class="search-box" placeholder="Search invoices..." id="searchInput" onkeyup="filterTable()">
            </div>
            <div class="table-wrapper">
                <table id="invoicesTable">
                    <thead>
                        <tr>
                            <th onclick="sortTable(0)">Date</th>
                            <th onclick="sortTable(1)">Sender</th>
                            <th onclick="sortTable(2)">Subject</th>
                            <th onclick="sortTable(3)">Amount</th>
                            <th onclick="sortTable(4)">Currency</th>
                            <th>Evidence</th>
                        </tr>
                    </thead>
                    <tbody>
                        {''.join(f"""
                        <tr>
                            <td>{row.get('date_utc', '')[:10]}</td>
                            <td class="truncate" title="{row.get('sender', '')}">{row.get('sender', '')[:30]}</td>
                            <td class="truncate" title="{row.get('subject', '')}">{row.get('subject', '')[:40]}</td>
                            <td class="amount">{CURRENCY_SYMBOLS.get(row.get('currency', 'UNK'), '')}{float(row.get('amount', 0)):,.2f}</td>
                            <td><span class="currency-badge {row.get('currency', 'UNK')}">{row.get('currency', 'UNK')}</span></td>
                            <td class="truncate" title="{row.get('evidence', '').replace('"', '&quot;')}">{row.get('evidence', '')[:50]}</td>
                        </tr>
                        """ for row in sorted(file_rows, key=lambda x: x.get('date_utc', ''), reverse=True)) if file_rows else '<tr><td colspan="6" class="empty-state">No invoices found</td></tr>'}
                    </tbody>
                </table>
            </div>
        </div>

        <div class="table-section">
            <div class="table-header">
                <h3 class="table-title">By Sender ({len(sender_totals)} senders)</h3>
            </div>
            <div class="table-wrapper">
                <table>
                    <thead>
                        <tr>
                            <th>Sender</th>
                            <th>Currency</th>
                            <th>Total Amount</th>
                        </tr>
                    </thead>
                    <tbody>
                        {''.join(f"""
                        <tr>
                            <td class="truncate" title="{sender}">{sender[:40]}</td>
                            <td><span class="currency-badge {curr}">{curr}</span></td>
                            <td class="amount">{CURRENCY_SYMBOLS.get(curr, '')}{amt:,.2f}</td>
                        </tr>
                        """ for (sender, curr), amt in sorted(total_by_sender_currency.items(), key=lambda x: -x[1]))}
                    </tbody>
                </table>
            </div>
        </div>
    </div>

    <script>
        const currencyData = {json.dumps(currency_data)};
        const topSenders = {json.dumps([{"sender": s[:25], "amount": a} for s, a in top_senders])};
        const monthlyData = {json.dumps({m: dict(d) for m, d in monthly_totals.items()})};
        const allRows = {json.dumps(file_rows)};

        // Currency Chart
        if (currencyData.length > 0) {{
            new Chart(document.getElementById('currencyChart'), {{
                type: 'doughnut',
                data: {{
                    labels: currencyData.map(d => d.currency),
                    datasets: [{{
                        data: currencyData.map(d => d.amount),
                        backgroundColor: ['#22c55e', '#3b82f6', '#f59e0b', '#ec4899', '#a0a0b0'],
                        borderWidth: 0
                    }}]
                }},
                options: {{
                    responsive: true,
                    maintainAspectRatio: false,
                    plugins: {{
                        legend: {{
                            position: 'right',
                            labels: {{ color: '#a0a0b0' }}
                        }}
                    }}
                }}
            }});
        }}

        // Senders Chart
        if (topSenders.length > 0) {{
            new Chart(document.getElementById('sendersChart'), {{
                type: 'bar',
                data: {{
                    labels: topSenders.map(d => d.sender),
                    datasets: [{{
                        label: 'Total Amount',
                        data: topSenders.map(d => d.amount),
                        backgroundColor: '#6366f1',
                        borderRadius: 4
                    }}]
                }},
                options: {{
                    indexAxis: 'y',
                    responsive: true,
                    maintainAspectRatio: false,
                    plugins: {{
                        legend: {{ display: false }}
                    }},
                    scales: {{
                        x: {{
                            grid: {{ color: '#2a2a3a' }},
                            ticks: {{ color: '#a0a0b0' }}
                        }},
                        y: {{
                            grid: {{ display: false }},
                            ticks: {{ color: '#a0a0b0' }}
                        }}
                    }}
                }}
            }});
        }}

        // Monthly Chart
        const months = Object.keys(monthlyData).sort();
        if (months.length > 0) {{
            const currencies = [...new Set(Object.values(monthlyData).flatMap(d => Object.keys(d)))];
            const colorMap = {{ ILS: '#22c55e', USD: '#3b82f6', EUR: '#f59e0b', GBP: '#ec4899', UNK: '#a0a0b0' }};

            new Chart(document.getElementById('monthlyChart'), {{
                type: 'bar',
                data: {{
                    labels: months,
                    datasets: currencies.map(curr => ({{
                        label: curr,
                        data: months.map(m => monthlyData[m]?.[curr] || 0),
                        backgroundColor: colorMap[curr] || '#6366f1',
                        borderRadius: 2
                    }}))
                }},
                options: {{
                    responsive: true,
                    maintainAspectRatio: false,
                    plugins: {{
                        legend: {{
                            labels: {{ color: '#a0a0b0' }}
                        }}
                    }},
                    scales: {{
                        x: {{
                            stacked: true,
                            grid: {{ color: '#2a2a3a' }},
                            ticks: {{ color: '#a0a0b0' }}
                        }},
                        y: {{
                            stacked: true,
                            grid: {{ color: '#2a2a3a' }},
                            ticks: {{ color: '#a0a0b0' }}
                        }}
                    }}
                }}
            }});
        }}

        function filterTable() {{
            const input = document.getElementById('searchInput').value.toLowerCase();
            const rows = document.querySelectorAll('#invoicesTable tbody tr');
            rows.forEach(row => {{
                const text = row.textContent.toLowerCase();
                row.style.display = text.includes(input) ? '' : 'none';
            }});
        }}

        let sortDir = {{}};
        function sortTable(col) {{
            const table = document.getElementById('invoicesTable');
            const tbody = table.querySelector('tbody');
            const rows = Array.from(tbody.querySelectorAll('tr'));

            sortDir[col] = !sortDir[col];

            rows.sort((a, b) => {{
                let aVal = a.cells[col]?.textContent || '';
                let bVal = b.cells[col]?.textContent || '';

                if (col === 3) {{
                    aVal = parseFloat(aVal.replace(/[^0-9.-]/g, '')) || 0;
                    bVal = parseFloat(bVal.replace(/[^0-9.-]/g, '')) || 0;
                    return sortDir[col] ? aVal - bVal : bVal - aVal;
                }}

                return sortDir[col] ? aVal.localeCompare(bVal) : bVal.localeCompare(aVal);
            }});

            rows.forEach(row => tbody.appendChild(row));
        }}

        function exportCSV() {{
            const headers = ['Date', 'Sender', 'Subject', 'Amount', 'Currency', 'Evidence'];
            const csvContent = [
                headers.join(','),
                ...allRows.map(row => [
                    row.date_utc || '',
                    '"' + (row.sender || '').replace(/"/g, '""') + '"',
                    '"' + (row.subject || '').replace(/"/g, '""') + '"',
                    row.amount || '',
                    row.currency || '',
                    '"' + (row.evidence || '').replace(/"/g, '""') + '"'
                ].join(','))
            ].join('\\n');

            const blob = new Blob([csvContent], {{ type: 'text/csv;charset=utf-8;' }});
            const link = document.createElement('a');
            link.href = URL.createObjectURL(blob);
            link.download = 'invoices_export.csv';
            link.click();
        }}
    </script>
</body>
</html>'''

    return html


def create_zip_archive(out_dir: Path, account_label: str, logger: logging.Logger, perf: PerformanceTracker) -> Path:
    """Create a zip file containing all outputs."""
    logger.info(f"📦 Creating zip archive...")
    zip_path = out_dir / f"{sanitize_filename(account_label)}_expenses_export.zip"

    with perf.timer("create_zip"):
        file_count = 0
        total_size = 0

        with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zf:
            for root, dirs, files in os.walk(out_dir):
                for file in files:
                    file_path = Path(root) / file
                    if file_path == zip_path:
                        continue
                    arcname = file_path.relative_to(out_dir)
                    zf.write(file_path, arcname)
                    file_count += 1
                    total_size += file_path.stat().st_size

        zip_size = zip_path.stat().st_size
        logger.info(f"   Archived {file_count} files ({format_bytes(total_size)} -> {format_bytes(zip_size)})")

    return zip_path


# ============== MAIN PROCESSING ==============

def run_account(
    account_label: str,
    creds_path: Path,
    token_path: Path,
    out_dir: Path,
    keywords: List[str],
    after_yyyy_mm_dd: str,
    max_messages: int,
    max_attachment_mb: int,
    enable_ocr: bool,
    ocr_max_pages: int,
    allow_duplicates: bool,
    create_zip: bool,
    logger: logging.Logger,
    perf: PerformanceTracker,
) -> None:
    logger.info("")
    logger.info("=" * 60)
    logger.info(f"🚀 STARTING ACCOUNT: {account_label}")
    logger.info("=" * 60)
    logger.info("")

    # Auth
    creds = load_or_auth(creds_path, token_path, logger, perf)
    service = gmail_service(creds)

    # Build query
    query = build_gmail_query(keywords, after_yyyy_mm_dd)
    logger.info(f"🔍 Search parameters:")
    logger.info(f"   Keywords: {len(keywords)} terms")
    logger.info(f"   Date filter: after {after_yyyy_mm_dd}")
    logger.info(f"   Max messages: {max_messages}")
    logger.info(f"   OCR enabled: {enable_ocr}")
    logger.debug(f"   Full query: {query}")

    # Get messages
    msg_ids = list_messages(service, query, max_messages, logger, perf)

    if not msg_ids:
        logger.warning("No messages found matching criteria!")
        return

    # Setup directories
    ensure_dir(out_dir)
    downloads_dir = out_dir / "downloads"
    ensure_dir(downloads_dir)

    log_file = out_dir / "processing.log"

    # Add file handler for this account
    file_handler = logging.FileHandler(log_file, encoding='utf-8')
    file_handler.setLevel(logging.DEBUG)
    file_handler.setFormatter(logging.Formatter(
        '%(asctime)s | %(levelname)-8s | %(message)s',
        datefmt='%Y-%m-%d %H:%M:%S'
    ))
    logger.addHandler(file_handler)

    results_path = out_dir / "results.jsonl"
    totals_csv = out_dir / "expenses_totals.csv"
    by_sender_csv = out_dir / "expenses_by_sender.csv"
    by_currency_csv = out_dir / "expenses_by_currency.csv"
    dashboard_path = out_dir / "dashboard.html"

    total_by_currency = defaultdict(float)
    total_by_sender_currency = defaultdict(float)
    file_rows = []
    total_bytes_downloaded = 0

    seen_hashes = set()

    logger.info("")
    logger.info(f"📧 Processing {len(msg_ids)} messages...")
    logger.info("")

    with results_path.open("w", encoding="utf-8") as f_out:
        for i, msg_id in enumerate(msg_ids, 1):
            perf.increment("messages_processed")

            # Progress logging
            if i == 1 or i % 25 == 0 or i == len(msg_ids):
                perf.log_progress(i, len(msg_ids), "messages")

            logger.debug(f"")
            logger.debug(f"--- Message {i}/{len(msg_ids)}: {msg_id} ---")

            with perf.timer("fetch_message"):
                try:
                    msg = service.users().messages().get(userId="me", id=msg_id, format="full").execute()
                except HttpError as e:
                    logger.error(f"   Failed to fetch message {msg_id}: {e}")
                    perf.increment("messages_failed")
                    continue

            payload = msg.get("payload", {}) or {}
            headers = payload.get("headers", []) or []

            from_raw = get_header(headers, "From")
            sender_name, sender_email = parseaddr(from_raw)
            sender_email = (sender_email or "").lower().strip()
            sender_key = sender_email or sanitize_filename(from_raw or "unknown_sender")

            subject = get_header(headers, "Subject")
            date_utc = parse_gmail_internal_date(msg.get("internalDate", ""))

            logger.debug(f"   From: {sender_email or from_raw[:50]}")
            logger.debug(f"   Subject: {subject[:60]}...")
            logger.debug(f"   Date: {date_utc[:10] if date_utc else 'unknown'}")

            plain, html = extract_text_parts(payload)
            links = list(dict.fromkeys(extract_links(plain) + extract_links(html)))

            attachments_meta = iter_attachments(payload)
            logger.debug(f"   Attachments found: {len(attachments_meta)}")

            downloaded = []
            analyzed = []

            sender_folder = downloads_dir / sanitize_filename(sender_key)
            ensure_dir(sender_folder)

            for att in attachments_meta:
                filename = att["filename"]
                ext = Path(filename).suffix.lower()

                if ext not in ALLOWED_EXTS:
                    logger.debug(f"      Skipping {filename} (unsupported extension: {ext})")
                    perf.increment("attachments_skipped_ext")
                    continue

                size = int(att.get("size") or 0)
                if size > max_attachment_mb * 1024 * 1024:
                    logger.debug(f"      Skipping {filename} (too large: {format_bytes(size)})")
                    perf.increment("attachments_skipped_size")
                    continue

                safe_name = sanitize_filename(filename)
                target = sender_folder / f"{msg_id}_{safe_name}"

                try:
                    logger.debug(f"      Downloading: {filename} ({format_bytes(size)})")

                    with perf.timer("download_attachment"):
                        data = download_attachment(service, msg_id, att["attachmentId"])

                    h = sha256_bytes(data)
                    if (not allow_duplicates) and (h in seen_hashes):
                        logger.debug(f"      Skipping duplicate (sha256: {h[:16]}...)")
                        perf.increment("attachments_skipped_duplicate")
                        continue
                    seen_hashes.add(h)

                    target.write_bytes(data)
                    total_bytes_downloaded += len(data)
                    perf.increment("attachments_downloaded")
                    perf.increment("bytes_downloaded", len(data))

                    logger.debug(f"      Saved: {target.relative_to(out_dir)}")

                    downloaded.append({
                        "filename": filename,
                        "saved_as": str(target.relative_to(out_dir)),
                        "mimeType": att.get("mimeType", ""),
                        "size": len(data),
                        "sha256": h,
                    })

                    # Analyze the file
                    logger.debug(f"      Analyzing content...")
                    analysis = analyze_file(target, enable_ocr, ocr_max_pages, logger, perf)
                    best = analysis["best_total"]

                    analyzed_item = {
                        "file": str(target.relative_to(out_dir)),
                        "analysis": analysis,
                    }
                    analyzed.append(analyzed_item)

                    if best:
                        amt = float(best["amount"])
                        curr = best["currency"]
                        total_by_currency[curr] += amt
                        total_by_sender_currency[(sender_key, curr)] += amt
                        perf.increment("invoices_detected")

                        file_rows.append({
                            "account": account_label,
                            "date_utc": date_utc,
                            "sender": sender_key,
                            "subject": subject,
                            "file": str(target.relative_to(out_dir)),
                            "currency": curr,
                            "amount": f"{amt:.2f}",
                            "evidence": best["context"].replace("\n", " ")[:180],
                            "method": best["source"],
                        })

                        logger.info(f"   💰 Found: {CURRENCY_SYMBOLS.get(curr, '')}{amt:,.2f} {curr} from {sender_key[:30]}")

                except HttpError as e:
                    logger.error(f"      Download failed: {e}")
                    perf.increment("attachments_failed")
                except Exception as e:
                    logger.error(f"      Processing failed: {e}")
                    perf.increment("attachments_failed")

            record = {
                "account": account_label,
                "message_id": msg_id,
                "date_utc": date_utc,
                "from_raw": from_raw,
                "sender_email": sender_email,
                "sender_name": sender_name,
                "sender_key": sender_key,
                "subject": subject,
                "links": links,
                "attachments_downloaded": downloaded,
                "attachments_analyzed": analyzed,
            }
            f_out.write(json.dumps(record, ensure_ascii=False) + "\n")

    # Write CSVs
    logger.info("")
    logger.info("📝 Writing output files...")

    with perf.timer("write_csvs"):
        with totals_csv.open("w", encoding="utf-8", newline="") as f:
            w = csv.DictWriter(f, fieldnames=[
                "account", "date_utc", "sender", "subject", "file",
                "currency", "amount", "evidence", "method"
            ])
            w.writeheader()
            for r in file_rows:
                w.writerow(r)
        logger.info(f"   ✅ {totals_csv.name}: {len(file_rows)} rows")

        with by_sender_csv.open("w", encoding="utf-8", newline="") as f:
            w = csv.writer(f)
            w.writerow(["sender", "currency", "sum_amount"])
            sender_rows = 0
            for (sender, curr), amt in sorted(total_by_sender_currency.items(), key=lambda x: (-x[1], x[0][0], x[0][1])):
                w.writerow([sender, curr, f"{amt:.2f}"])
                sender_rows += 1
        logger.info(f"   ✅ {by_sender_csv.name}: {sender_rows} rows")

        with by_currency_csv.open("w", encoding="utf-8", newline="") as f:
            w = csv.writer(f)
            w.writerow(["currency", "sum_amount"])
            for curr, amt in sorted(total_by_currency.items(), key=lambda x: -x[1]):
                w.writerow([curr, f"{amt:.2f}"])
        logger.info(f"   ✅ {by_currency_csv.name}: {len(total_by_currency)} rows")

    # Generate dashboard
    logger.info("   Generating dashboard...")

    with perf.timer("generate_dashboard"):
        generated_at = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

        # Prepare perf summary for dashboard
        total_time = time.time() - perf.start_time
        perf_summary = {
            "messages_processed": perf.get_count("messages_processed"),
            "attachments_downloaded": perf.get_count("attachments_downloaded"),
            "total_bytes": format_bytes(perf.get_count("bytes_downloaded")),
            "ocr_performed": perf.get_count("ocr_attempts"),
            "processing_time": perf._format_duration(total_time),
        }

        dashboard_html = generate_dashboard_html(
            account_label,
            file_rows,
            dict(total_by_currency),
            dict(total_by_sender_currency),
            generated_at,
            perf_summary
        )
        dashboard_path.write_text(dashboard_html, encoding="utf-8")
    logger.info(f"   ✅ {dashboard_path.name}")

    # Create zip archive
    if create_zip:
        zip_path = create_zip_archive(out_dir, account_label, logger, perf)
        logger.info(f"   ✅ {zip_path.name}")

    # Summary
    logger.info("")
    logger.info("=" * 60)
    logger.info(f"📊 ACCOUNT SUMMARY: {account_label}")
    logger.info("=" * 60)
    logger.info(f"   Messages processed: {perf.get_count('messages_processed')}")
    logger.info(f"   Attachments downloaded: {perf.get_count('attachments_downloaded')}")
    logger.info(f"   Data downloaded: {format_bytes(total_bytes_downloaded)}")
    logger.info(f"   Invoices detected: {perf.get_count('invoices_detected')}")
    logger.info("")
    logger.info("   💰 TOTALS BY CURRENCY:")
    for curr, amt in sorted(total_by_currency.items(), key=lambda x: -x[1]):
        logger.info(f"      {curr}: {CURRENCY_SYMBOLS.get(curr, '')}{amt:,.2f}")
    logger.info("")
    logger.info(f"   📁 Output directory: {out_dir}")
    logger.info("=" * 60)

    # Remove file handler
    logger.removeHandler(file_handler)
    file_handler.close()


def main():
    p = argparse.ArgumentParser(
        description="Invoice expense tracker with dashboard & zip export (Gmail).",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  python invoice_expenses.py --accounts personal --ocr
  python invoice_expenses.py --accounts work personal --ocr -v
  python invoice_expenses.py --accounts business --max 1000 --no-zip

Setup:
  1. Go to https://console.cloud.google.com/
  2. Create a project and enable Gmail API
  3. Create OAuth 2.0 credentials (Desktop app)
  4. Download credentials.json to current directory
  5. Run script - browser will open for auth
        """
    )
    p.add_argument("--credentials", default="credentials.json", help="Path to OAuth credentials.json")
    p.add_argument("--out", default="out", help="Output directory")
    p.add_argument("--accounts", nargs="+", required=True, help="Labels for accounts (authorize each separately)")
    p.add_argument("--max", type=int, default=4000, help="Max messages per account")
    p.add_argument("--max-attachment-mb", type=int, default=25, help="Skip attachments bigger than this")
    p.add_argument("--keywords", nargs="*", default=None, help="Override keywords list")
    p.add_argument("--ocr", action="store_true", help="Enable OCR for PDFs/images")
    p.add_argument("--ocr-max-pages", type=int, default=3, help="OCR first N pages of PDFs")
    p.add_argument("--allow-duplicates", action="store_true", help="Count identical files multiple times")
    p.add_argument("--no-zip", action="store_true", help="Skip creating zip archive")
    p.add_argument("-v", "--verbose", action="store_true", help="Enable verbose/debug logging")
    args = p.parse_args()

    # Setup logging
    base_out = Path(args.out).expanduser().resolve()
    ensure_dir(base_out)

    logger = setup_logging(verbose=args.verbose, log_file=base_out / "main.log")
    perf = PerformanceTracker(logger)

    logger.info("")
    logger.info("╔════════════════════════════════════════════════════════════╗")
    logger.info("║        INVOICE EXPENSE TRACKER WITH DASHBOARD              ║")
    logger.info("║              Gmail API + OCR + Analytics                   ║")
    logger.info("╚════════════════════════════════════════════════════════════╝")
    logger.info("")

    creds_path = Path(args.credentials).expanduser().resolve()
    if not creds_path.exists():
        logger.error(f"❌ Missing credentials file: {creds_path}")
        logger.error("")
        logger.error("To get credentials.json:")
        logger.error("  1. Go to https://console.cloud.google.com/")
        logger.error("  2. Create a new project (or select existing)")
        logger.error("  3. Enable the Gmail API:")
        logger.error("     - Go to 'APIs & Services' > 'Library'")
        logger.error("     - Search for 'Gmail API' and enable it")
        logger.error("  4. Create OAuth credentials:")
        logger.error("     - Go to 'APIs & Services' > 'Credentials'")
        logger.error("     - Click 'Create Credentials' > 'OAuth client ID'")
        logger.error("     - Choose 'Desktop app' as application type")
        logger.error("     - Download the JSON file")
        logger.error("  5. Rename the downloaded file to 'credentials.json'")
        logger.error("  6. Place it in the current directory or specify --credentials path")
        logger.error("")
        sys.exit(1)

    keywords = args.keywords if args.keywords else DEFAULT_KEYWORDS

    after_date = date.today() - timedelta(days=365)
    after_str = after_date.strftime("%Y/%m/%d")

    logger.info(f"📋 Configuration:")
    logger.info(f"   Accounts: {', '.join(args.accounts)}")
    logger.info(f"   Date range: {after_str} to today")
    logger.info(f"   Max messages per account: {args.max}")
    logger.info(f"   OCR enabled: {args.ocr}")
    logger.info(f"   Verbose mode: {args.verbose}")
    logger.info(f"   Output directory: {base_out}")
    logger.info("")

    for label in args.accounts:
        token_path = base_out / "tokens" / f"token_{sanitize_filename(label)}.json"
        out_dir = base_out / sanitize_filename(label)

        try:
            run_account(
                account_label=label,
                creds_path=creds_path,
                token_path=token_path,
                out_dir=out_dir,
                keywords=keywords,
                after_yyyy_mm_dd=after_str,
                max_messages=args.max,
                max_attachment_mb=args.max_attachment_mb,
                enable_ocr=args.ocr,
                ocr_max_pages=args.ocr_max_pages,
                allow_duplicates=args.allow_duplicates,
                create_zip=not args.no_zip,
                logger=logger,
                perf=perf,
            )
        except Exception as e:
            logger.error(f"❌ Failed processing account '{label}': {e}")
            if args.verbose:
                import traceback
                logger.error(traceback.format_exc())

    # Final summary
    perf.print_summary()

    logger.info("")
    logger.info("✅ All done! Open dashboard.html in your browser to view results.")
    logger.info("")


if __name__ == "__main__":
    main()
